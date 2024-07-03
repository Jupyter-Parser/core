import builtins
from typing import List
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shape import InlineShape
from docx.table import Table
from docx.text.paragraph import Run, Paragraph

from .utils.table_of_content import add_toc
from .utils.add_link import add_hyperlink
from .types.Elements import Element, MdDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm
from . import types as tp
from .types import Cell, CellType, Section
from io import BytesIO
import requests
from haggis.files.docx import list_number


def scale_picture(picture: InlineShape, new_width):
    aspect_ratio = float(picture.height) / float(picture.width)
    picture.width = new_width
    picture.height = int(aspect_ratio * new_width)


def generate_document(
    cells: List[Cell],
    generate_toc: bool = False,
    cwd: str = "",
    section_opts: Section = Section()
) -> bytes:
    document = Document()

    max_picture_width = section_opts.width - \
        (section_opts.left + section_opts.right)

    section = document.sections[0]

    section.page_width = Mm(section_opts.width)
    section.page_height = Mm(section_opts.height)
    section.top_margin = Mm(section_opts.top)
    section.right_margin = Mm(section_opts.right)
    section.bottom_margin = Mm(section_opts.bottom)
    section.left_margin = Mm(section_opts.left)

    styles = document.styles

    paragraph_style = styles.add_style("Параграф", WD_STYLE_TYPE.PARAGRAPH)
    latex_style = styles.add_style("Latex", WD_STYLE_TYPE.CHARACTER)
    quote_style = styles.add_style("Цитата", WD_STYLE_TYPE.PARAGRAPH)
    code_style = styles.add_style("Код", WD_STYLE_TYPE.PARAGRAPH)
    character_code_style = styles.add_style(
        "Код (символы)", WD_STYLE_TYPE.CHARACTER)
    code_output_style = styles.add_style("Вывод", WD_STYLE_TYPE.PARAGRAPH)
    toc_style = styles.add_style("Оглавление", WD_STYLE_TYPE.PARAGRAPH)

    for cell in cells:
        if cell.cell_type == CellType.code:
            code = document.add_paragraph()
            code.style = code_style
            code.add_run(cell.source)

            for output in cell.outputs:
                code_output = document.add_paragraph()
                code_output.style = code_output_style
                run = code_output.add_run()

                if output.text:
                    try:
                        run.add_text(output.text)
                        run.add_break()
                    except:
                        pass
                if output.image:
                    picture = run.add_picture(BytesIO(output.image))
                    if picture.width > max_picture_width:
                        scale_picture(picture, max_picture_width)
        elif cell.cell_type == CellType.markdown:
            root: MdDocument = cell.source

            def add_text(children: List[Element], paragraph: Paragraph, italic=False, bold=False):
                for element in children:
                    if type(element) == str:
                        run: Run = paragraph.add_run()
                        run.bold = bold
                        run.italic = italic
                        run.add_text(element)
                    elif type(element) == tp.MdCode:
                        run: Run = paragraph.add_run()
                        run.bold = bold
                        run.italic = italic
                        run.style = character_code_style
                        run.add_text(element.children)
                    else:
                        walk(element)

            def add_runs(
                children: List[Element], paragraph: Paragraph, level=0, prev=None
            ):
                for element in children:
                    match type(element):
                        case builtins.str:
                            paragraph.add_run(element)
                        case tp.MdEmphasis:
                            add_text(element.children, paragraph, italic=True)
                        case tp.MdStrongEmphasis:
                            add_text(element.children, paragraph, bold=True)
                        case tp.MdLatex:
                            run: Run = paragraph.add_run()
                            run.style = latex_style
                            run.add_text(element.latex)
                        case tp.MdCode:
                            run: Run = paragraph.add_run()
                            run.style = character_code_style
                            run.add_text(element.children)
                        case tp.MdLink:
                            run: Run = paragraph.add_run()

                            txt = ""

                            for child in children:
                                if type(child) == str:
                                    txt += child

                            add_hyperlink(run, txt, element.source)
                        case _:
                            walk(element, level, prev)

            def walk(element: Element, level=0, prev=None):
                if not element:
                    return
                match type(element):
                    case tp.MdHeading:
                        level = element.level
                        h = document.add_heading(level=level)
                        add_runs(element.children, h)
                    case tp.MdLine:
                        document.add_paragraph()
                    case tp.MdDocument:
                        for child in element.children:
                            walk(child)
                    case tp.MdParagraph:
                        p = document.add_paragraph()
                        p.style = paragraph_style
                        add_runs(element.children, p)
                    case tp.MdImage:
                        try:
                            src = element.source

                            if "http://" in src or "https://" in src:
                                img = requests.get(src).content
                            else:
                                img = open(f"{cwd}/{src}", "rb").read()

                            picture = document.add_picture(BytesIO(img))
                            if picture.width > max_picture_width:
                                scale_picture(picture, max_picture_width)
                        except:
                            pass
                    case tp.MdQuote:
                        p = document.add_paragraph()
                        p.style = quote_style

                        for child in element.children:
                            if type(child) == tp.MdParagraph:
                                add_runs(child.children, p)

                    case tp.MdList:
                        level += 1
                        ordered = element.ordered
                        start = element.start

                        if level == 1:
                            style = "List Number" if ordered else "List Bullet"
                        else:
                            style = (
                                f"List Number {level}"
                                if ordered
                                else f"List Bullet {level}"
                            )

                        for item in element.children:
                            p = document.add_paragraph(style=style)

                            for child in item.children:
                                if type(child) == tp.MdParagraph:
                                    add_runs(child.children, p, level, prev)
                                elif type(child) == tp.MdList:
                                    walk(child, level, prev)

                            list_number(document, p, num=start)

                            start += 1
                    case tp.MdTable:
                        rows: List[tp.MdTableRow] = element.children

                        table: Table = document.add_table(
                            rows=0, cols=len(rows[0].children)
                        )

                        table.style = "TableGrid"

                        for r in rows:
                            row = table.add_row().cells

                            for i, c in enumerate(r.children):
                                p = row[i].add_paragraph()

                                match c.align:
                                    case "left":
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                    case "right":
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                    case "center":
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                add_runs(c.children, p)
                    case _:
                        p = document.add_paragraph()

                        add_runs(element.children, p)

            walk(root)

    if generate_toc:
        add_toc(
            document,
            "Оглавление",
            heading_style=toc_style,
            levels=(1, 6)
        )

    file = BytesIO()

    document.save(file)

    file.seek(0)

    return file.read()
