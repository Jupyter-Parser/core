from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.styles.style import ParagraphStyle


def add_toc(
    document: DocumentObject,
    heading: str | None = None,
    heading_level: int | None = 1,
    heading_style: ParagraphStyle | None = None,
    levels: tuple[int, int] = (1, 3),
):
    document.add_page_break()

    if heading:
        h = document.add_heading(heading, heading_level)

        if heading_style:
            h.style = heading_style

    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f'TOC \\o "{levels[0]}-{levels[1]}" \\h \\z \\u'

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:updateFields")
    fldChar3.set(qn("w:val"), "true")

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
